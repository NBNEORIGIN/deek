"""
ingest_manuals.py — bulk-ingest NBNE machinery manuals into Deek memory.

Walks a folder of manuals (default: D:\\Google Drive\\My Drive\\001 NBNE\\002 BLANKS),
parses each file by type, chunks the text, embeds, and upserts into
``claw_code_chunks`` with ``chunk_type='manual'`` so the agent's
``search_manuals`` tool can find them.

File types handled:
    .pdf            — pypdf text extraction
    .docx           — python-docx
    .txt / .md      — utf-8 decode
    .png / .jpg / .heic / etc. — Claude vision OCR (anthropic SDK direct)

Machine name encoding:
    NBNE refers to its kit by nicknames (Hulk, Beast, Rolf, Mao, …). The
    folder convention is ONE subdirectory per machine — e.g.
        BLANKS/Hulk/operator-manual.pdf
        BLANKS/Beast/maintenance-record-2025-q1.pdf
    The script derives machine name from the parent folder. Override
    with --machine for ad-hoc ingests of files outside that layout.

    The machine name is encoded into ``chunk_name`` as
        ``"<machine> · <filename> · chunk-<n>"``
    so search_manuals's optional ``machine`` filter can ILIKE-prefix-match
    on chunk_name without needing a schema migration. (When the manual
    library grows large enough that ILIKE filtering hurts, we'll add a
    dedicated machine_name column — until then this keeps things simple.)

Usage:
    # Dry-run a single folder, see what it would do
    python -m scripts.ingest_manuals \\
        --folder "D:/Google Drive/My Drive/001 NBNE/002 BLANKS" \\
        --dry-run

    # Actually write
    python -m scripts.ingest_manuals \\
        --folder "D:/Google Drive/My Drive/001 NBNE/002 BLANKS" \\
        --commit

    # Ingest one machine, override the auto-derived name
    python -m scripts.ingest_manuals \\
        --folder "D:/scratch/manuals" \\
        --machine "Hulk" \\
        --commit

Hard caps:
    * 50 MB per file — skip with a warning above this. Stops a stray
      30-page scanned-photo PDF from drowning the embedding budget.
    * 200 chunks per file — each chunk ~2000 chars. A 200-page manual
      maps to roughly 400-600 chunks at typical density; we cap at
      200 to keep one document from dominating retrieval. (Operator
      manuals tend to be more relevant than service manuals; the
      first 200 chunks usually catch operations.)
"""
from __future__ import annotations

import argparse
import hashlib
import io
import logging
import os
import sys
from pathlib import Path
from typing import Any

# Resolve repo root so this works whether invoked as a module or a file.
REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

try:
    from dotenv import load_dotenv
    load_dotenv(REPO_ROOT / '.env')
except ImportError:
    pass

logger = logging.getLogger('ingest_manuals')

MAX_FILE_BYTES = 50 * 1024 * 1024
MAX_CHUNKS_PER_FILE = 200
CHUNK_TARGET_CHARS = 2000
CHUNK_OVERLAP_CHARS = 200

TEXT_EXTS = {'.txt', '.md', '.log'}
PDF_EXTS = {'.pdf'}
DOCX_EXTS = {'.docx'}
IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.heic', '.bmp'}


def _slugify_machine(name: str) -> str:
    """Light-touch normalisation — strip whitespace, preserve case so
    "Hulk" still reads naturally in the chunk_name. ILIKE handles
    case-insensitive matching on the search side."""
    return (name or '').strip().strip('"').strip("'") or 'unknown'


def _derive_machine_from_path(file_path: Path, root: Path) -> str:
    """Walk up from the file and return the first directory under
    ``root`` — that's the machine folder. Falls back to '(uncategorised)'
    for files dropped directly in the root."""
    try:
        rel = file_path.resolve().relative_to(root.resolve())
    except ValueError:
        return '(uncategorised)'
    parts = rel.parts
    if len(parts) <= 1:
        return '(uncategorised)'
    return _slugify_machine(parts[0])


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    out: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ''
        except Exception as exc:
            text = f'[page {i + 1} unreadable: {type(exc).__name__}]'
        if text.strip():
            out.append(f'--- page {i + 1} ---\n{text.strip()}')
    return '\n\n'.join(out).strip()


def _extract_docx(path: Path) -> str:
    import docx as docx_mod
    doc = docx_mod.Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        line = (p.text or '').strip()
        if line:
            parts.append(line)
    for ti, table in enumerate(doc.tables):
        parts.append(f'\n[table {ti + 1}]')
        for row in table.rows:
            cells = [(c.text or '').strip() for c in row.cells]
            parts.append(' | '.join(cells))
    return '\n'.join(parts).strip()


def _extract_text(path: Path) -> str:
    return path.read_text(encoding='utf-8', errors='replace')


def _extract_image_via_claude(path: Path) -> str:
    """OCR + caption a manual photo via Claude vision.

    NBNE often has paper-only documents — operator quick-cards taped to
    a machine, hand-written maintenance notes in a logbook, schematic
    photos. Claude vision handles all three reasonably well in one
    pass: returns text content + structural description. We embed that
    description as the chunk content; the agent then retrieves a
    natural-language summary of the photo when relevant.
    """
    import base64
    try:
        from anthropic import Anthropic
    except ImportError as exc:
        raise RuntimeError(
            'anthropic SDK required for image OCR — pip install anthropic'
        ) from exc

    api_key = os.getenv('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        raise RuntimeError(
            'ANTHROPIC_API_KEY not set — image OCR unavailable. Either '
            'set the key or skip image files (--skip-images).'
        )

    blob = path.read_bytes()
    ext = path.suffix.lower()
    media_type = {
        '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.gif': 'image/gif', '.webp': 'image/webp',
    }.get(ext)
    if media_type is None:
        # HEIC/BMP — convert via Pillow first.
        from PIL import Image as PILImage
        if ext == '.heic':
            try:
                from pillow_heif import register_heif_opener
                register_heif_opener()
            except ImportError as exc:
                raise RuntimeError(
                    'HEIC support requires pillow-heif (pip install pillow-heif)'
                ) from exc
        img = PILImage.open(io.BytesIO(blob))
        if img.mode in ('RGBA', 'LA', 'P'):
            bg = PILImage.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        out_buf = io.BytesIO()
        img.save(out_buf, format='JPEG', quality=85, optimize=True)
        blob = out_buf.getvalue()
        media_type = 'image/jpeg'

    encoded = base64.b64encode(blob).decode('ascii')

    client = Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=os.getenv('CLAUDE_MODEL', 'claude-sonnet-4-6'),
        max_tokens=2000,
        messages=[{
            'role': 'user',
            'content': [
                {
                    'type': 'image',
                    'source': {
                        'type': 'base64',
                        'media_type': media_type,
                        'data': encoded,
                    },
                },
                {
                    'type': 'text',
                    'text': (
                        'This is a photograph or scan of a machinery manual, '
                        'maintenance record, control panel, schematic, or '
                        'paper document for NBNE\'s production equipment. '
                        'Transcribe ALL visible text verbatim, including part '
                        'numbers, model numbers, torque specs, voltages, dates, '
                        'serial numbers, and hand-written notes. Then add a '
                        'brief structural caption describing the visual content '
                        '(e.g. "diagram showing belt tensioner with arrows '
                        'indicating direction of rotation"). Output as plain '
                        'text — no markdown, no commentary. If the image is '
                        'illegible or contains nothing useful, write exactly '
                        '"[image: no recoverable content]".'
                    ),
                },
            ],
        }],
    )
    return ''.join(
        b.text for b in msg.content if getattr(b, 'type', '') == 'text'
    ).strip()


def _extract(path: Path, ext: str, *, skip_images: bool = False) -> str:
    if ext in PDF_EXTS:
        return _extract_pdf(path)
    if ext in DOCX_EXTS:
        return _extract_docx(path)
    if ext in TEXT_EXTS:
        return _extract_text(path)
    if ext in IMAGE_EXTS:
        if skip_images:
            return ''
        return _extract_image_via_claude(path)
    return ''


def _chunk(text: str, target: int = CHUNK_TARGET_CHARS,
           overlap: int = CHUNK_OVERLAP_CHARS) -> list[str]:
    """Sliding-window chunker. Prefers paragraph boundaries within the
    target size; falls back to mid-sentence splits if a paragraph
    exceeds the target."""
    text = (text or '').strip()
    if not text:
        return []
    if len(text) <= target:
        return [text]

    chunks: list[str] = []
    pos = 0
    n = len(text)
    while pos < n:
        end = min(pos + target, n)
        if end < n:
            # Try to break on a paragraph boundary, then sentence.
            for needle in ('\n\n', '. ', '\n'):
                cut = text.rfind(needle, pos, end)
                if cut > pos + (target // 2):  # don't break too early
                    end = cut + len(needle)
                    break
        chunk = text[pos:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        pos = max(pos + 1, end - overlap)
    return chunks


def _embed_fn():
    """Return the embed callable, or None if no embedder is configured."""
    try:
        from core.wiki.embeddings import get_embed_fn
        return get_embed_fn()
    except Exception as exc:
        logger.warning('embed_fn import failed: %s', exc)
        return None


def _connect_db():
    import psycopg2
    db_url = os.getenv('DATABASE_URL', '')
    if not db_url:
        raise RuntimeError('DATABASE_URL not set')
    return psycopg2.connect(db_url, connect_timeout=10)


def _upsert_chunk(conn, *, file_path: str, chunk_content: str,
                  chunk_name: str, embedding: list[float] | None) -> None:
    """Upsert one (file_path, chunk_name) row into claw_code_chunks
    with chunk_type='manual'. Idempotent — re-running the script over
    the same files replaces the same rows.

    Mirrors the wiki_tools._embed_into_chunks pattern (delete-then-insert
    keyed by project_id + file_path + chunk_type + chunk_name)."""
    content_hash = hashlib.sha256(chunk_content.encode('utf-8')).hexdigest()
    with conn.cursor() as cur:
        cur.execute(
            '''DELETE FROM claw_code_chunks
                WHERE project_id = 'deek'
                  AND file_path = %s
                  AND chunk_type = 'manual'
                  AND chunk_name = %s''',
            (file_path, chunk_name),
        )
        if embedding is None:
            cur.execute(
                '''INSERT INTO claw_code_chunks
                    (project_id, file_path, chunk_content, chunk_type,
                     chunk_name, content_hash, embedding, indexed_at,
                     salience, salience_signals,
                     last_accessed_at, access_count)
                   VALUES ('deek', %s, %s, 'manual', %s, %s, NULL,
                           NOW(), 5.0,
                           '{"via": "ingest_manuals"}'::jsonb,
                           NOW(), 0)''',
                (file_path, chunk_content, chunk_name, content_hash),
            )
        else:
            cur.execute(
                '''INSERT INTO claw_code_chunks
                    (project_id, file_path, chunk_content, chunk_type,
                     chunk_name, content_hash, embedding, indexed_at,
                     salience, salience_signals,
                     last_accessed_at, access_count)
                   VALUES ('deek', %s, %s, 'manual', %s, %s, %s::vector,
                           NOW(), 5.0,
                           '{"via": "ingest_manuals"}'::jsonb,
                           NOW(), 0)''',
                (file_path, chunk_content, chunk_name,
                 content_hash, embedding),
            )


def _process_file(
    path: Path, root: Path, machine_override: str | None,
    *, skip_images: bool, dry_run: bool, conn, embed_fn,
) -> dict:
    """Returns a stat dict for the report at the end."""
    stat = {
        'path': str(path),
        'machine': '',
        'ext': path.suffix.lower(),
        'size': 0,
        'chunks': 0,
        'embedded': 0,
        'skipped': '',
        'error': '',
    }
    try:
        stat['size'] = path.stat().st_size
        if stat['size'] > MAX_FILE_BYTES:
            stat['skipped'] = f'over {MAX_FILE_BYTES // (1024 * 1024)}MB'
            return stat

        ext = path.suffix.lower()
        if ext not in (TEXT_EXTS | PDF_EXTS | DOCX_EXTS | IMAGE_EXTS):
            stat['skipped'] = 'unsupported extension'
            return stat

        machine = machine_override or _derive_machine_from_path(path, root)
        stat['machine'] = machine

        text = _extract(path, ext, skip_images=skip_images)
        if not text.strip():
            stat['skipped'] = 'no extractable text'
            return stat

        chunks = _chunk(text)
        if len(chunks) > MAX_CHUNKS_PER_FILE:
            chunks = chunks[:MAX_CHUNKS_PER_FILE]
            stat['skipped'] = (
                f'truncated to first {MAX_CHUNKS_PER_FILE} chunks'
            )
        stat['chunks'] = len(chunks)

        if dry_run:
            return stat

        rel_path = str(path.relative_to(root)) if path.is_relative_to(root) else str(path)
        for i, chunk_text in enumerate(chunks):
            chunk_name = f'{machine} · {path.stem} · chunk-{i + 1:03d}'
            embedding: list[float] | None = None
            if embed_fn is not None:
                try:
                    # Prepend a context header before embedding so the
                    # query-time semantic match knows which machine
                    # the chunk belongs to even when the chunk content
                    # itself doesn't mention the nickname.
                    embedding = embed_fn(
                        f'Machine: {machine}\nManual: {path.name}\n\n{chunk_text[:6000]}'
                    )
                except Exception as exc:
                    logger.warning('embed failed on %s chunk %d: %s', path.name, i, exc)
            _upsert_chunk(
                conn,
                file_path=rel_path,
                chunk_content=chunk_text,
                chunk_name=chunk_name,
                embedding=embedding,
            )
            if embedding is not None:
                stat['embedded'] += 1
        conn.commit()
        return stat
    except Exception as exc:
        stat['error'] = f'{type(exc).__name__}: {exc}'
        try:
            conn.rollback()
        except Exception:
            pass
        return stat


def main(argv: list[str] | None = None) -> int:
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s %(levelname)s %(name)s — %(message)s',
        handlers=[logging.StreamHandler(sys.stdout)],
    )

    parser = argparse.ArgumentParser(
        prog='python -m scripts.ingest_manuals',
        description=__doc__.strip().split('\n')[0],
    )
    parser.add_argument(
        '--folder', required=True,
        help='Root folder of manuals. Subdirectories are treated as machine names.',
    )
    parser.add_argument(
        '--machine', default=None,
        help='Override machine name. Default: derived from parent folder.',
    )
    parser.add_argument(
        '--dry-run', action='store_true',
        help='Walk + extract + chunk but do not write to DB.',
    )
    parser.add_argument(
        '--commit', action='store_true',
        help='Actually write to DB. Required to land any rows.',
    )
    parser.add_argument(
        '--skip-images', action='store_true',
        help='Skip .png/.jpg/.heic etc — useful for a fast PDF-only pass '
             'when ANTHROPIC_API_KEY is unavailable or you want to avoid '
             'vision-API spend.',
    )
    args = parser.parse_args(argv)

    if args.dry_run and args.commit:
        print('--dry-run and --commit are mutually exclusive', file=sys.stderr)
        return 2
    dry_run = not args.commit

    root = Path(args.folder).expanduser().resolve()
    if not root.is_dir():
        print(f'folder not found or not a directory: {root}', file=sys.stderr)
        return 2

    logger.info('root: %s (dry_run=%s skip_images=%s)', root, dry_run, args.skip_images)

    conn = None
    embed_fn = None
    if not dry_run:
        try:
            conn = _connect_db()
            try:
                from pgvector.psycopg2 import register_vector
                register_vector(conn)
            except Exception:
                pass
        except Exception as exc:
            print(f'DB connection failed: {exc}', file=sys.stderr)
            return 1
        embed_fn = _embed_fn()
        if embed_fn is None:
            logger.warning(
                'no embedder available — chunks will be inserted without '
                'embeddings. Lexical (ILIKE) search will still work; '
                'semantic retrieval will not.'
            )

    files = sorted(p for p in root.rglob('*') if p.is_file())
    logger.info('found %d files', len(files))

    stats: list[dict] = []
    for path in files:
        stat = _process_file(
            path, root, args.machine,
            skip_images=args.skip_images,
            dry_run=dry_run,
            conn=conn,
            embed_fn=embed_fn,
        )
        stats.append(stat)
        tag = (
            stat['error'] or stat['skipped']
            or f"{stat['chunks']} chunks ({stat['embedded']} embedded)"
        )
        logger.info('  [%s] %s — %s', stat['machine'] or '?', path.name, tag)

    if conn is not None:
        try:
            conn.close()
        except Exception:
            pass

    # Summary
    n_total = len(stats)
    n_ingested = sum(1 for s in stats if s['chunks'] and not s['error'] and not s['skipped'])
    n_skipped = sum(1 for s in stats if s['skipped'])
    n_errored = sum(1 for s in stats if s['error'])
    chunks_total = sum(s['chunks'] for s in stats)
    embedded_total = sum(s['embedded'] for s in stats)

    by_machine: dict[str, int] = {}
    for s in stats:
        if s['machine']:
            by_machine[s['machine']] = by_machine.get(s['machine'], 0) + s['chunks']

    print(file=sys.stderr)
    print('───── ingest_manuals summary ─────', file=sys.stderr)
    print(f'  files seen:         {n_total}', file=sys.stderr)
    print(f'  files ingested:     {n_ingested}', file=sys.stderr)
    print(f'  files skipped:      {n_skipped}', file=sys.stderr)
    print(f'  files errored:      {n_errored}', file=sys.stderr)
    print(f'  total chunks:       {chunks_total}', file=sys.stderr)
    print(f'  total embeddings:   {embedded_total}', file=sys.stderr)
    if by_machine:
        print(f'  per-machine:', file=sys.stderr)
        for m, n in sorted(by_machine.items(), key=lambda kv: -kv[1]):
            print(f'    {m:<25} {n:>5} chunks', file=sys.stderr)
    if dry_run:
        print('  (dry-run — no rows written. Re-run with --commit to land.)',
              file=sys.stderr)

    return 0 if n_errored == 0 else 3


if __name__ == '__main__':
    sys.exit(main())

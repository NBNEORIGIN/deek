"""Voice surface — file upload + text extraction.

Receives one or more files via multipart/form-data and returns the
extracted text per file. The Next.js /voice page uploads files
attached to a chat message; the extracted text is then prepended to
the user's prompt before forwarding to the agent stream.

No file is persisted by default — extraction is in-memory and the
blob is discarded once the response is sent. If the request includes
``persist=true``, the original bytes are written to the volume at
``/app/data/uploads/<sha8>/<filename>`` for later retrieval (defer
this hook until we actually need it).

Supported formats:
    .pdf            — pypdf
    .docx           — python-docx
    .csv / .tsv     — stdlib csv (capped at 1k rows + summary)
    .xlsx / .xlsm   — openpyxl (first sheet, capped at 1k rows)
    .txt / .md      — plain decode
    .png/.jpg/.heic — accepted but flagged as not-yet-supported

Per-file size cap: 10 MB. Per-request: 5 files. Hard caps so the
agent's context isn't drowned by a stray 50MB report.
"""
from __future__ import annotations

import csv
import hashlib
import io
import logging
from typing import Any

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse

from api.middleware.auth import verify_api_key

log = logging.getLogger(__name__)
router = APIRouter(prefix='/voice', tags=['Voice file upload'])


MAX_FILE_BYTES = 10 * 1024 * 1024     # 10 MB
MAX_FILES_PER_REQUEST = 5
MAX_TEXT_CHARS = 60_000               # ~15k tokens per file
MAX_CSV_ROWS = 1_000
MAX_XLSX_ROWS = 1_000


def _sha8(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:8]


def _ext(filename: str) -> str:
    return ('.' + (filename or '').rsplit('.', 1)[-1].lower()) if '.' in (filename or '') else ''


def _truncate(s: str, limit: int = MAX_TEXT_CHARS) -> tuple[str, bool]:
    if len(s) <= limit:
        return s, False
    return s[:limit] + '\n…[truncated]', True


def _extract_pdf(blob: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(blob))
    out: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ''
        except Exception as exc:
            text = f'[page {i + 1} unreadable: {type(exc).__name__}]'
        out.append(f'--- page {i + 1} ---\n{text.strip()}')
    return '\n\n'.join(out).strip()


def _extract_docx(blob: bytes) -> str:
    import docx as docx_mod  # python-docx
    doc = docx_mod.Document(io.BytesIO(blob))
    parts: list[str] = []
    for p in doc.paragraphs:
        line = (p.text or '').strip()
        if line:
            parts.append(line)
    # Tables
    for ti, table in enumerate(doc.tables):
        parts.append(f'\n[table {ti + 1}]')
        for row in table.rows:
            cells = [(c.text or '').strip() for c in row.cells]
            parts.append(' | '.join(cells))
    return '\n'.join(parts).strip()


def _extract_csv(blob: bytes, sep: str = ',') -> str:
    text = blob.decode('utf-8', errors='replace')
    reader = csv.reader(io.StringIO(text), delimiter=sep)
    rows: list[list[str]] = []
    for i, row in enumerate(reader):
        if i >= MAX_CSV_ROWS:
            rows.append(['…', f'(truncated at {MAX_CSV_ROWS} rows)'])
            break
        rows.append(row)
    if not rows:
        return '[empty CSV]'
    header = rows[0]
    body = rows[1:]
    summary = (
        f'columns: {len(header)} | rows: {len(body)} | '
        f'header: {", ".join(header[:10])}{"…" if len(header) > 10 else ""}'
    )
    sample_lines = [' | '.join(r) for r in rows[:50]]
    return summary + '\n\n' + '\n'.join(sample_lines)


def _extract_xlsx(blob: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(filename=io.BytesIO(blob), read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        return '[empty workbook]'
    rows: list[list[str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i >= MAX_XLSX_ROWS:
            rows.append(['…', f'(truncated at {MAX_XLSX_ROWS} rows)'])
            break
        rows.append([('' if c is None else str(c)) for c in row])
    wb.close()
    if not rows:
        return '[empty sheet]'
    header = rows[0]
    body = rows[1:]
    summary = (
        f'sheet: {ws.title} | columns: {len(header)} | rows: {len(body)} | '
        f'header: {", ".join(header[:10])}{"…" if len(header) > 10 else ""}'
    )
    sample_lines = [' | '.join(r) for r in rows[:50]]
    return summary + '\n\n' + '\n'.join(sample_lines)


def _extract_text(blob: bytes) -> str:
    return blob.decode('utf-8', errors='replace')


_EXTRACTORS: dict[str, Any] = {
    '.pdf': _extract_pdf,
    '.docx': _extract_docx,
    '.csv': lambda b: _extract_csv(b, ','),
    '.tsv': lambda b: _extract_csv(b, '\t'),
    '.xlsx': _extract_xlsx,
    '.xlsm': _extract_xlsx,
    '.txt': _extract_text,
    '.md': _extract_text,
    '.log': _extract_text,
    '.json': _extract_text,
}

_IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.heic', '.gif', '.webp', '.bmp'}


@router.post('/upload')
async def voice_upload(
    files: list[UploadFile] = File(...),
    _: bool = Depends(verify_api_key),
) -> JSONResponse:
    """Extract text from one or more uploaded files.

    Returns:
        {
          "files": [
            {
              "name": "report.pdf",
              "size": 123456,
              "type": ".pdf",
              "sha8": "a1b2c3d4",
              "text": "...extracted...",
              "truncated": false,
              "supported": true,
              "error": null
            },
            ...
          ]
        }
    """
    if not files:
        raise HTTPException(status_code=400, detail='no files')
    if len(files) > MAX_FILES_PER_REQUEST:
        raise HTTPException(
            status_code=400,
            detail=f'too many files (max {MAX_FILES_PER_REQUEST})',
        )

    results: list[dict] = []
    for upload in files:
        name = upload.filename or 'upload'
        ext = _ext(name)
        size = 0
        text = ''
        truncated = False
        supported = ext in _EXTRACTORS or ext in _IMAGE_EXTS
        error: str | None = None

        try:
            blob = await upload.read()
            size = len(blob)
            if size > MAX_FILE_BYTES:
                error = f'file exceeds {MAX_FILE_BYTES // (1024 * 1024)} MB limit'
                supported = False
            elif ext in _IMAGE_EXTS:
                # Phase 2 — vision model integration not wired yet.
                error = (
                    'image upload received but vision support is not yet '
                    'wired into the chat path; this attachment will be '
                    'acknowledged but not analysed'
                )
                supported = False
            elif ext in _EXTRACTORS:
                extractor = _EXTRACTORS[ext]
                raw = extractor(blob)
                text, truncated = _truncate(raw or '')
            else:
                error = f'unsupported extension: {ext or "(none)"}'
                supported = False
        except Exception as exc:
            error = f'{type(exc).__name__}: {exc}'
            log.warning('[voice/upload] %s failed: %s', name, exc)
        finally:
            try:
                await upload.close()
            except Exception:
                pass

        results.append({
            'name': name,
            'size': size,
            'type': ext,
            'sha8': _sha8(blob) if size else '',
            'text': text,
            'truncated': truncated,
            'supported': supported,
            'error': error,
        })

    return JSONResponse({'files': results})
